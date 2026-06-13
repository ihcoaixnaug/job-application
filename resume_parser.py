"""Parse Word (.docx) and PDF resumes into plain text."""
import re
from pathlib import Path


def _clean_text(text: str) -> str:
    """Replace runs of whitespace/tabs with a single space, strip lines."""
    text = re.sub(r"[ \t]+", " ", text)
    lines = [ln.strip() for ln in text.splitlines()]
    # Drop duplicate consecutive lines (merged-cell artefacts)
    deduped: list[str] = []
    for ln in lines:
        if ln and (not deduped or ln != deduped[-1]):
            deduped.append(ln)
    return "\n".join(deduped)


def parse_docx(filepath: str | Path) -> str:
    """Extract all text from a .docx file, including tables."""
    from docx import Document  # lazy import — python-docx
    doc = Document(str(filepath))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = re.sub(r"[ \t]+", " ", para.text).strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_cells = [re.sub(r"[ \t]+", " ", cell.text).strip()
                         for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append(" | ".join(row_cells))

    return _clean_text("\n".join(parts))


def parse_pdf(filepath: str | Path) -> str:
    """Extract all text from a PDF file."""
    import pypdf
    reader = pypdf.PdfReader(str(filepath))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def parse_resume(filepath: str | Path) -> str:
    """Auto-detect format and parse .docx or .pdf."""
    p = Path(filepath)
    if p.suffix.lower() == ".pdf":
        return parse_pdf(p)
    return parse_docx(p)
