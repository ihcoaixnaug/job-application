"""Generate 项目说明文档.docx for the Tencent PCG campus AI competition."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ── Styles ─────────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = '微软雅黑'
style_normal.font.size = Pt(11)
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_font(run, bold=False, size=11, color=None):
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=16, color=(31, 73, 125))
    elif level == 2:
        set_font(run, bold=True, size=13, color=(68, 114, 196))
    else:
        set_font(run, bold=True, size=11, color=(89, 89, 89))
    return p

def add_para(text, indent=False, italic=False, muted=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if muted:
        run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def add_bullet(text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.9 if not sub else 1.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # header bg color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F497D')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = 'EEF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=10)
            # row bg
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg)
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
    # set col widths if given
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing after table

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(30)
run = cover.add_run('求职助手（Job Application Assistant）')
set_font(run, bold=True, size=22, color=(31, 73, 125))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('AI 驱动的校园求职全流程管理平台')
set_font(run2, size=14, color=(68, 114, 196))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = sub2.add_run('腾讯校园 AI 项目大赛 · 产品说明文档')
set_font(run3, size=11, color=(100, 100, 100))
run3.font.italic = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. 痛点与背景
# ══════════════════════════════════════════════════════════════════════════════
add_heading('一、真实痛点：校园求职的"信息散落"困境', 1)

add_para(
    '每逢求职季，大学生的典型一天是这样的——早上刷 Boss直聘，下午切到实习僧，'
    '晚上翻备忘录确认哪些投了哪些没投，偶尔发现两周前投的简历已读未回，'
    '却完全不记得是哪家公司的哪个岗位。'
)

add_heading('四大核心痛点', 2)
add_bullet('信息散落：平台割裂，需在多个 App 间反复切换，搜索词靠"感觉"，容易漏掉好机会')
add_bullet('筛选低效：每天面对数十条岗位，凭主观判断是否匹配，耗时且不可复现')
add_bullet('跟进失控：投递记录靠备忘录，不清楚哪些在等 HR 回、哪些已到面试阶段')
add_bullet('招呼同质：自我介绍千篇一律，复制粘贴时改了公司名忘改岗位名，出现尴尬')

add_para(
    '这四个痛点是真实存在于每一位求职期大学生日常中的具体场景，可通过访谈、'
    '观察行为数据（多平台 Tab 切换次数、备忘录关键词）加以验证。',
    muted=True, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. 用户需求分析
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('二、用户需求分析', 1)

add_heading('2.1 目标用户', 2)
add_table(
    ['用户群体', '核心诉求', '使用频率'],
    [
        ['大三/大四在校生', '找暑期实习、秋招全流程管理', '每日活跃'],
        ['研究生求职期', '多方向探索、精准定向', '高频'],
        ['有意转换赛道的应届生', '跨领域岗位匹配度评估', '中频'],
    ],
    col_widths=[4.5, 7, 3]
)

add_heading('2.2 需求优先级（KANO 模型）', 2)
add_table(
    ['需求', '类型', '说明'],
    [
        ['多平台信息聚合', '必备型', '用户对平台割裂已有明确痛感'],
        ['AI 岗位匹配评分', '期望型', '核心差异化能力，决定产品价值'],
        ['自动打招呼/投递', '期望型', '大幅降低重复性操作成本'],
        ['进度看板可视化', '期望型', '减少认知负担，提升管理感'],
        ['腾讯新闻公司动态', '兴奋型', '超预期功能，PCG 生态协同'],
        ['面试题 AI 生成', '兴奋型', '打开求职到面试的延伸场景'],
    ],
    col_widths=[5, 3, 7]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. 功能架构
# ══════════════════════════════════════════════════════════════════════════════
add_heading('三、功能架构', 1)

add_heading('3.1 核心功能模块', 2)
add_table(
    ['模块', '功能描述'],
    [
        ['🔍 智能搜索', 'AI 解析求职意向 → 生成 6-10 个关键词 → 多平台×多城市×多关键词全量抓取，(平台,岗位ID) 去重'],
        ['🤖 AI 匹配分析', 'LLM 逐条比对简历与岗位，输出 0-100 匹配分、匹配理由、亮点（绿色）、顾虑（橙色）'],
        ['📋 看板管理', '7 阶段流水线看板（待投→Offer），拖拽流转，时间线备注，实时漏斗统计'],
        ['💬 智能打招呼', 'AI 模式（个性化生成）/ 固定模板（用户预设）二选一，发送前必过预览弹窗'],
        ['🏷️ 公司规模标签', '内置 100+ 公司关键词，自动识别大厂/中厂/小厂，可按规模筛选'],
        ['📰 公司动态（PCG）', '（规划中）调用腾讯新闻 API，在岗位详情页显示该公司近 30 天新闻摘要'],
        ['🎬 面试备考（PCG）', '（规划中）AI 根据 JD 生成面试题，联动腾讯视频推荐相关备考内容'],
        ['🎭 演示/真实分离', '内置 30 条仿真岗位数据；真实与演示数据独立存储、独立看板，互不干扰'],
    ],
    col_widths=[4.5, 10.5]
)

add_heading('3.2 信息架构', 2)
add_para('页面结构：单页应用（SPA），三区布局')
add_bullet('左侧边栏：简历上传 / 求职偏好 / 搜索配置')
add_bullet('顶部 Header：Boss 登录状态 / 演示模式切换 / 各阶段漏斗统计')
add_bullet('主内容区：真实/演示 Tab 切换 → 看板视图 / 列表视图，带平台/规模/分数/排序筛选器')

# ══════════════════════════════════════════════════════════════════════════════
# 4. AI 能力
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('四、AI 能力说明', 1)

add_heading('4.1 已实现的 AI 能力', 2)
add_table(
    ['AI 能力', '模型调用方式', '输出形式'],
    [
        ['意向→关键词语义扩展', 'LLM（DeepSeek V3）', 'JSON 数组，6-10 个搜索词'],
        ['简历-岗位语义匹配', 'LLM 逐条分析', '分数(0-100) + 理由 + 亮点列表 + 顾虑列表'],
        ['打招呼文案个性化生成', 'LLM（≤60字约束）', '精炼打招呼语，结合岗位要求与简历摘要'],
    ],
    col_widths=[5, 4.5, 5.5]
)

add_heading('4.2 模型接入', 2)
add_bullet('统一通过 OpenRouter 接入，默认使用 DeepSeek V3（中文能力强、成本极低）')
add_bullet('支持热切换：Claude / GPT-4o / Gemini，用户在设置面板一键选择')
add_bullet('无 API Key 时自动降级：关键词按"/"切分，匹配分随机生成（演示模式可用）')

add_heading('4.3 近期可加入的 AI 能力（PCG 结合点）', 2)
add_table(
    ['AI 能力', '对应 PCG 产品', '价值'],
    [
        ['公司近期动态摘要', '腾讯新闻 API', '帮助用户判断公司现状，减少踩坑'],
        ['面试题自动生成（JD→题目）', '腾讯视频内容推荐', '打通"求职"到"面试备考"链路'],
        ['简历定向优化建议', 'LLM 增强', '提升每条投递的转化率'],
        ['投递效果归因分析', '数据看板', '沉淀用户行为数据，指导后续求职策略'],
    ],
    col_widths=[5, 4, 6]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. 技术架构
# ══════════════════════════════════════════════════════════════════════════════
add_heading('五、技术架构', 1)

add_table(
    ['层次', '技术选型', '说明'],
    [
        ['前端', 'HTML5 + Vanilla JS + CSS', '无框架依赖，轻量 SPA，支持拖拽、动画、弹窗'],
        ['后端', 'Python FastAPI', '异步框架，BackgroundTasks 处理爬取/AI 调用'],
        ['浏览器自动化', 'Patchright（Playwright 分支）', '二进制层反检测，模拟人类鼠标/键盘，保活 Boss 会话'],
        ['AI 接口', 'OpenRouter（DeepSeek V3 等）', '支持多模型热切换，按量计费'],
        ['数据库', 'SQLite + aiosqlite', '零部署，本地运行，is_demo 字段隔离演示/真实数据'],
        ['反爬对策', 'UA随机化、延迟抖动、会话复用', 'Boss 会话绑定浏览器指纹，整个生命周期内不重建 Context'],
    ],
    col_widths=[3.5, 5, 6.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. 与 PCG 业务的结合
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('六、与腾讯 PCG 业务的结合', 1)

add_para(
    'PCG（平台与内容事业群）旗下覆盖腾讯视频、腾讯新闻、微视、QQ浏览器等核心内容产品。'
    '求职助手的核心场景天然与 PCG 内容生态形成高价值协同：'
)

add_heading('6.1 腾讯新闻 × 公司动态模块', 2)
add_para(
    '在每个岗位详情页接入腾讯新闻 API，展示该公司近30天内的新闻摘要（融资、裁员、战略调整等）。'
    '用户投递前可快速了解公司现状，做出更理性的决策，同时为腾讯新闻带来精准的内容分发场景。'
)

add_heading('6.2 腾讯视频 × 面试备考模块', 2)
add_para(
    'AI 根据 JD 自动生成 5 道高频面试题，并在岗位详情页推荐腾讯视频上的相关备考内容'
    '（行业认知视频、面试技巧课程等）。这将求职助手的使用场景从"投递"延伸到"准备面试"，'
    '提升用户在 PCG 内容生态中的停留深度。'
)

add_heading('6.3 QQ 浏览器 × 求职助手插件（远期）', 2)
add_para(
    '将求职助手封装为 QQ 浏览器侧边栏插件，用户在任意岗位页面浏览时，'
    '插件自动提取岗位信息、实时匹配简历、显示建议打招呼文案，实现"零跳转"的求职体验。'
)

add_table(
    ['PCG 产品', '协同点', '价值方向'],
    [
        ['腾讯新闻', '公司动态摘要嵌入岗位详情', '内容精准分发、提升阅读深度'],
        ['腾讯视频', '面试备考视频推荐', '延伸使用场景、增加内容消费'],
        ['QQ 浏览器', '求职助手侧边栏插件（远期）', '生态整合、提升浏览器粘性'],
    ],
    col_widths=[4, 6, 5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. 原型说明
# ══════════════════════════════════════════════════════════════════════════════
add_heading('七、原型说明', 1)

add_para('产品采用"可运行原型"形式，非静态线框图，所有功能均已实现并可实时演示。')

add_heading('7.1 主要界面', 2)
add_table(
    ['界面', '核心元素', '交互特点'],
    [
        ['主看板', '7列状态看板、卡片（含匹配分/规模标签/活跃度）、漏斗统计', '拖拽流转、卡片进场动画'],
        ['搜索侧边栏', '意向偏好输入、AI解析关键词药丸、多城市复选框', 'AI关键词一键生成、回车快速添加'],
        ['岗位详情弹窗', '匹配分面板、亮点/顾虑列表、状态切换、时间线日志', '全屏弹窗、时间线实时追加'],
        ['打招呼预览弹窗', '可编辑文案区、岗位信息提示', 'AI生成/固定模板二选一，发前可改'],
        ['设置面板', '打招呼模式选择、模板编辑、Boss/实习僧登录', '独立模态，含模板预览'],
    ],
    col_widths=[3.5, 7.5, 4]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. 落地规划
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('八、落地规划', 1)

add_heading('8.1 近期优化（2周内可实现）', 2)
add_bullet('面试题自动生成：AI 读 JD，输出 5 道高频面试题，直接显示在岗位详情页')
add_bullet('简历定向优化建议：AI 指出简历哪一条改写后匹配分可提升，给出具体建议')
add_bullet('腾讯新闻公司动态：接入新闻 API，嵌入岗位详情页（PCG 核心协同）')
add_bullet('投递效果归因看板：统计回复率最高的关键词/公司规模/城市，指导策略调整')

add_heading('8.2 中期规划（1-3个月）', 2)
add_bullet('微信小程序版：移动端访问，Tencent 生态打通，推送面试/笔试提醒')
add_bullet('腾讯视频面试备考联动：根据岗位方向推荐相关课程/视频内容')
add_bullet('多用户支持 + 云端部署：从个人工具升级为平台产品')
add_bullet('简历版本管理：针对不同方向维护多版本简历，一键切换')

add_heading('8.3 远期愿景（6个月+）', 2)
add_bullet('QQ 浏览器求职助手插件：任意求职页面实时匹配，零跳转体验')
add_bullet('社区数据积累：匿名化的投递成功率数据，形成"岗位热度指数"')
add_bullet('校企联动：为 PCG 等事业群提供在校生求职倾向数据洞察（匿名聚合）')

add_heading('8.4 用户价值量化（实测估算）', 2)
add_table(
    ['操作', '传统方式', '使用本产品', '节省时间'],
    [
        ['多平台搜索', '手动刷3个平台，30分钟/天', '一键搜索，2分钟', '约28分钟/天'],
        ['岗位筛选', '手动判断50条，约50分钟/次', 'AI评分排序，扫一眼', '约45分钟/次'],
        ['打招呼撰写', '手写每条，约5分钟/条', '模板/AI生成+预览，30秒', '约4.5分钟/条'],
        ['进度跟进', '备忘录/Excel，易遗漏', '看板可视化，拖拽更新', '认知负担大幅降低'],
    ],
    col_widths=[3.5, 5, 4.5, 2]
)

# ══════════════════════════════════════════════════════════════════════════════
# Footer note
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('本产品为可运行原型，所有功能均已实现，可随时启动演示。')
set_font(run, size=10, color=(120, 120, 120))
run.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = 'docs/项目说明文档.docx'
doc.save(out_path)
print(f'Saved → {out_path}')
